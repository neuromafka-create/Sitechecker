# auth/rec_report.py — генерация .docx документа с рекомендациями по аутентификации

import re
from datetime import datetime
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches, Cm


COLOR_DARK    = RGBColor(0x26, 0x32, 0x38)
COLOR_BLUE    = RGBColor(0x15, 0x65, 0xC0)
COLOR_RED     = RGBColor(0xC6, 0x28, 0x28)
COLOR_ORANGE  = RGBColor(0xE6, 0x5C, 0x00)
COLOR_GREEN   = RGBColor(0x2E, 0x7D, 0x32)
COLOR_GRAY    = RGBColor(0x55, 0x55, 0x55)


def _set_paragraph_border_bottom(paragraph, color="2E75B6", size=6):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _shade_paragraph(paragraph, fill_hex: str):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def _heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size  = Pt(15)
        run.font.color.rgb = COLOR_DARK
        _set_paragraph_border_bottom(p, "37474F", 8)
    elif level == 2:
        run.font.size  = Pt(13)
        run.font.color.rgb = COLOR_BLUE
    else:
        run.font.size  = Pt(11)
        run.font.color.rgb = COLOR_DARK
    return p


def _code_block(doc: Document, code: str):
    for line in code.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Inches(0.3)
        _shade_paragraph(p, "F5F5F5")
        run = p.add_run(line if line else " ")
        run.font.name  = "Courier New"
        run.font.size  = Pt(9)
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)


def _add_inline_formatting(paragraph, text: str):
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = COLOR_DARK
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name  = "Courier New"
            run.font.size  = Pt(10)
            run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
        elif part:
            run = paragraph.add_run(part)
            run.font.size  = Pt(11)
            run.font.color.rgb = COLOR_DARK


def _parse_and_render(doc: Document, text: str):
    lines = text.split("\n")
    in_code = False
    code_buf = []

    for line in lines:
        stripped = line.rstrip()
        if stripped.startswith("```"):
            if in_code:
                _code_block(doc, "\n".join(code_buf))
                code_buf = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_buf.append(line)
            continue
        if stripped.startswith("## "):
            _heading(doc, stripped[3:].strip(), level=1)
        elif stripped.startswith("### "):
            _heading(doc, stripped[4:].strip(), level=2)
        elif stripped.startswith("#### "):
            _heading(doc, stripped[5:].strip(), level=3)
        elif re.match(r"^(СРОЧНО|ВАЖНО|РЕКОМЕНДУЕТСЯ|КРИТИЧНО|ВЫСОКИЙ|СРЕДНИЙ)\b", stripped, re.I):
            m = re.match(r"^(\S+)\b[:\s]*(.*)", stripped)
            if m:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after  = Pt(2)
                badge = p.add_run(f" {m.group(1).upper()} ")
                badge.bold = True
                badge.font.size = Pt(10)
                if "СРОЧНО" in m.group(1) or "КРИТИЧНО" in m.group(1):
                    badge.font.color.rgb = COLOR_RED
                elif "ВАЖНО" in m.group(1) or "ВЫСОКИЙ" in m.group(1):
                    badge.font.color.rgb = COLOR_ORANGE
                else:
                    badge.font.color.rgb = COLOR_GREEN
                desc = p.add_run(f"  {m.group(2).strip()}")
                desc.font.size = Pt(11)
                desc.font.color.rgb = COLOR_DARK
        elif stripped.startswith("---") or stripped.startswith("==="):
            p = doc.add_paragraph()
            _set_paragraph_border_bottom(p, "BDBDBD", 4)
        elif re.match(r"^\d+\.\s", stripped):
            text_part = re.sub(r"^\d+\.\s+", "", stripped)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            p.paragraph_format.left_indent  = Inches(0.3)
            _add_inline_formatting(p, text_part)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            p.paragraph_format.left_indent  = Inches(0.3)
            _add_inline_formatting(p, stripped[2:])
        elif stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            _add_inline_formatting(p, stripped)
        else:
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    if in_code and code_buf:
        _code_block(doc, "\n".join(code_buf))


def generate_auth_recommendations_docx(result: dict, recommendations_text: str) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)

    domain = result.get("domain", "сайт")
    url = result.get("url", "")
    risk_level = result.get("risk_level", "clean")
    hits = result.get("hits") or []
    pages_count = len(result.get("pages_checked") or [])
    date = datetime.now().strftime("%d.%m.%Y")

    risk_labels = {
        "critical": "🔴 КРИТИЧЕСКИЙ", "high": "🟠 ВЫСОКИЙ",
        "medium": "🟡 СРЕДНИЙ", "low": "🔵 ВНИМАНИЕ",
        "clean": "✅ Чисто", "error": "⚪ Ошибка",
    }

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(4)
    title_run = title_p.add_run("РЕКОМЕНДАЦИИ ПО АУТЕНТИФИКАЦИИ")
    title_run.bold = True
    title_run.font.size  = Pt(16)
    title_run.font.color.rgb = COLOR_DARK

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("проверка соответствия требованиям РФ")
    sub_run.font.size  = Pt(13)
    sub_run.font.color.rgb = COLOR_GRAY
    _set_paragraph_border_bottom(subtitle_p, "BDBDBD", 4)
    doc.add_paragraph()

    meta_p = doc.add_paragraph()
    _shade_paragraph(meta_p, "ECEFF1")
    for label, value in [
        ("Сайт", url or domain),
        ("Дата аудита", date),
        ("Оценка риска", risk_labels.get(risk_level, risk_level)),
        ("Найдено проблем", str(len(hits))),
        ("Страниц проверено", str(pages_count)),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Inches(0.2)
        _shade_paragraph(p, "ECEFF1")
        bold_r = p.add_run(f"{label}: ")
        bold_r.bold = True
        bold_r.font.size = Pt(11)
        bold_r.font.color.rgb = COLOR_DARK
        val_r = p.add_run(value)
        val_r.font.size = Pt(11)
        val_r.font.color.rgb = COLOR_GRAY
    doc.add_paragraph()

    if hits:
        _heading(doc, "Найденные нарушения", level=1)
        for h in hits:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Inches(0.3)
            sev = h.get("severity", "info")
            icon = {"critical": "🔴", "high": "🟠", "medium": "🟡", "low": "🔵"}.get(sev, "⚪")
            p.add_run(f"{icon}  ").font.color.rgb = {
                "critical": COLOR_RED, "high": COLOR_ORANGE,
            }.get(sev, COLOR_GRAY)
            p.add_run(h.get("name", "—")).font.size = Pt(11)
        doc.add_paragraph()

    _parse_and_render(doc, recommendations_text)

    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_border_bottom(footer_p, "BDBDBD", 2)
    footer_r = footer_p.add_run(
        f"Документ сформирован автоматически на основе проверки аутентификации  •  {date}"
    )
    footer_r.font.size  = Pt(9)
    footer_r.font.color.rgb = COLOR_GRAY

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
