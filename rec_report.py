# rec_report.py — генерация .docx документа с рекомендациями по результатам аудита
#
# Использует python-docx.
# Парсит markdown-подобный текст от HUBRIS и оформляет в Word-документ.

import re
from datetime import datetime
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches, Cm


# ── Цвета ─────────────────────────────────────────────────────────────────────
COLOR_DARK    = RGBColor(0x26, 0x32, 0x38)   # заголовки
COLOR_BLUE    = RGBColor(0x15, 0x65, 0xC0)   # разделы
COLOR_RED     = RGBColor(0xC6, 0x28, 0x28)   # нарушения / СРОЧНО
COLOR_ORANGE  = RGBColor(0xE6, 0x5C, 0x00)   # ВАЖНО
COLOR_GREEN   = RGBColor(0x2E, 0x7D, 0x32)   # OK
COLOR_GRAY    = RGBColor(0x55, 0x55, 0x55)   # вспомогательный текст


def _set_paragraph_border_bottom(paragraph, color="2E75B6", size=6):
    """Добавляет нижнюю границу параграфа (используется вместо горизонтальной линии)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _shade_paragraph(paragraph, fill_hex: str):
    """Заливка фона параграфа."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def _heading(doc: Document, text: str, level: int = 1):
    """Добавляет заголовок с нужным форматированием."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True

    if level == 1:
        run.font.size  = Pt(15)
        run.font.color.rgb = COLOR_DARK
        _set_paragraph_border_bottom(p, "37474F", 8)
    elif level == 2:
        run.font.size  = Pt(13)
        run.font.color.rgb = COLOR_BLUE
    else:
        run.font.size  = Pt(11)
        run.font.color.rgb = COLOR_DARK

    return p


def _body(doc: Document, text: str, indent: bool = False):
    """Обычный параграф."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.size  = Pt(11)
    run.font.color.rgb = COLOR_DARK
    return p


def _bullet(doc: Document, text: str, indent_level: int = 0):
    """Маркированный пункт."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Inches(0.3 + indent_level * 0.25)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_DARK
    return p


def _priority_badge(doc: Document, priority: str, description: str):
    """Строка с бейджем приоритета (СРОЧНО / ВАЖНО / РЕКОМЕНДУЕТСЯ)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)

    badge = p.add_run(f" {priority} ")
    badge.bold = True
    badge.font.size = Pt(10)
    if "СРОЧНО" in priority:
        badge.font.color.rgb = COLOR_RED
    elif "ВАЖНО" in priority:
        badge.font.color.rgb = COLOR_ORANGE
    else:
        badge.font.color.rgb = COLOR_GREEN

    space = p.add_run("  ")
    space.font.size = Pt(11)

    desc = p.add_run(description)
    desc.font.size = Pt(11)
    desc.font.color.rgb = COLOR_DARK
    return p


def _code_block(doc: Document, code: str):
    """Блок кода с моноширинным шрифтом и серым фоном."""
    # Разбиваем по строкам
    lines = code.strip().split("\n")
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Inches(0.3)
        _shade_paragraph(p, "F5F5F5")
        run = p.add_run(line if line else " ")
        run.font.name  = "Courier New"
        run.font.size  = Pt(9)
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)


# ── Парсер markdown-подобного текста HUBRIS ─────────────────────────────────

def _parse_and_render(doc: Document, text: str):
    """
    Разбирает текст от HUBRIS (markdown-подобный) и рендерит в документ.
    Обрабатывает: ## заголовки, ### подзаголовки, - списки, ```код```, **жирный**.
    """
    lines = text.split("\n")
    in_code = False
    code_buf = []

    for line in lines:
        stripped = line.rstrip()

        # ── Блоки кода ────────────────────────────────────────────────────
        if stripped.startswith("```"):
            if in_code:
                _code_block(doc, "\n".join(code_buf))
                code_buf = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_buf.append(line)
            continue

        # ── Заголовки ─────────────────────────────────────────────────────
        if stripped.startswith("## "):
            _heading(doc, stripped[3:].strip(), level=1)
        elif stripped.startswith("### "):
            _heading(doc, stripped[4:].strip(), level=2)
        elif stripped.startswith("#### "):
            _heading(doc, stripped[5:].strip(), level=3)

        # ── Приоритеты ────────────────────────────────────────────────────
        elif re.match(r"^(СРОЧНО|ВАЖНО|РЕКОМЕНДУЕТСЯ)\b", stripped, re.I):
            m = re.match(r"^(\S+)\b[:\s]*(.*)", stripped)
            if m:
                _priority_badge(doc, m.group(1).upper(), m.group(2).strip())

        # ── Горизонтальная линия ──────────────────────────────────────────
        elif stripped.startswith("---") or stripped.startswith("==="):
            p = doc.add_paragraph()
            _set_paragraph_border_bottom(p, "BDBDBD", 4)

        # ── Нумерованный список ───────────────────────────────────────────
        elif re.match(r"^\d+\.\s", stripped):
            text_part = re.sub(r"^\d+\.\s+", "", stripped)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            p.paragraph_format.left_indent  = Inches(0.3)
            _add_inline_formatting(p, text_part)

        # ── Маркированный список ──────────────────────────────────────────
        elif stripped.startswith("- ") or stripped.startswith("* "):
            text_part = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            p.paragraph_format.left_indent  = Inches(0.3)
            _add_inline_formatting(p, text_part)

        # ── Обычный текст ─────────────────────────────────────────────────
        elif stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            _add_inline_formatting(p, stripped)

        # ── Пустая строка ─────────────────────────────────────────────────
        else:
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Незакрытый блок кода
    if in_code and code_buf:
        _code_block(doc, "\n".join(code_buf))


def _add_inline_formatting(paragraph, text: str):
    """Обрабатывает **жирный** и `код` внутри текста."""
    # Разбиваем по **..** и `..`
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = COLOR_DARK
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name  = "Courier New"
            run.font.size  = Pt(10)
            run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
        elif part:
            run = paragraph.add_run(part)
            run.font.size  = Pt(11)
            run.font.color.rgb = COLOR_DARK


# ── Главная функция генерации .docx ──────────────────────────────────────────

def generate_recommendations_docx(result: dict, recommendations_text: str) -> bytes:
    """
    Генерирует .docx документ с рекомендациями.

    result                — словарь результатов аудита (из checker.py)
    recommendations_text  — текст от HUBRIS

    Возвращает bytes для отдачи через Flask send_file.
    """
    doc = Document()

    # ── Поля страницы ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Базовый шрифт документа
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)

    domain = result.get("domain", "сайт")
    risk   = result.get("risk", "—")
    date   = datetime.now().strftime("%d.%m.%Y")

    # ── Шапка документа ───────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(4)
    title_run = title_p.add_run("РЕКОМЕНДАЦИИ ПО УСТРАНЕНИЮ НАРУШЕНИЙ")
    title_run.bold = True
    title_run.font.size  = Pt(16)
    title_run.font.color.rgb = COLOR_DARK

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(2)
    sub_run = subtitle_p.add_run("требований 152-ФЗ «О персональных данных»")
    sub_run.font.size  = Pt(13)
    sub_run.font.color.rgb = COLOR_GRAY
    _set_paragraph_border_bottom(subtitle_p, "BDBDBD", 4)

    doc.add_paragraph()

    # ── Реквизиты аудита ──────────────────────────────────────────────────────
    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_after = Pt(2)
    _shade_paragraph(meta_p, "ECEFF1")

    meta_items = [
        ("Сайт",             result.get("url") or domain),
        ("Дата аудита",      date),
        ("Оценка риска",     risk),
        ("HTTP-статус",      str(result.get("http_status", "—"))),
        ("Нарушений",        str(len(result.get("violations") or []))),
    ]
    for label, value in meta_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Inches(0.2)
        _shade_paragraph(p, "ECEFF1")
        bold_r = p.add_run(f"{label}: ")
        bold_r.bold = True
        bold_r.font.size = Pt(11)
        bold_r.font.color.rgb = COLOR_DARK
        val_r = p.add_run(value)
        val_r.font.size = Pt(11)
        val_r.font.color.rgb = COLOR_GRAY

    doc.add_paragraph()

    # ── Сводка нарушений ──────────────────────────────────────────────────────
    violations = result.get("violations") or []
    if violations:
        _heading(doc, "Выявленные нарушения", level=1)
        for v in violations:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.left_indent  = Inches(0.3)
            # Иконка
            icon_r = p.add_run("✗  ")
            icon_r.bold = True
            icon_r.font.color.rgb = COLOR_RED
            icon_r.font.size = Pt(11)
            # Текст нарушения
            text_r = p.add_run(v)
            text_r.font.size = Pt(11)
            text_r.font.color.rgb = COLOR_DARK

        doc.add_paragraph()

    # ── Основной текст рекомендаций от HUBRIS ───────────────────────────────
    _parse_and_render(doc, recommendations_text)

    # ── Нижний колонтитул ─────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_border_bottom(footer_p, "BDBDBD", 2)
    footer_r = footer_p.add_run(
        f"Документ сформирован автоматически на основе аудита по 152-ФЗ  •  {date}"
    )
    footer_r.font.size  = Pt(9)
    footer_r.font.color.rgb = COLOR_GRAY

    # ── Сохраняем в bytes ─────────────────────────────────────────────────────
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
